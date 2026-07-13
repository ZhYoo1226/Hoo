from .base_state import BaseState,StateOwner
from pathlib import Path
from common import GlobalFunction
from .state_doc import WordParseState,HeadingChunk
from docx import Document
from docx.shared import Pt
import json,threading,re

# 正文字符数 ≤ 此阈值时跳过 LLM，直接用原文作摘要（~200 token，中文字符约 1:1.5~2）
SUMMARY_LLM_THRESHOLD_CHARS = 130

# 生成节点摘要
class GenSummaryState(BaseState):
    """对单个文本切块节点生成摘要，原地修改 chunk.prefix_summary / chunk.summary"""
    stateName = "GenSummaryState"

    def __init__(self, chunk: HeadingChunk, **kwargs):
        super().__init__(**kwargs)
        self.chunk = chunk
        self._done = False
        self._duration = 1000 * 60 * 8

    def Enter(self, owner):
        lines = []
        for child in self.chunk.children:
            if child.summary:
                lines.append(f'{child.number} {child.title}：{child.summary}')
        self.children_summaries = '\n'.join(lines)
        
        self.node_number = self.chunk.number
        self.node_title = self.chunk.title
        self.node_level = self.chunk.level
        self.node_content = self.chunk.content

        # 正文足够短 → 跳过 LLM，直接用原文当摘要（小文本跳过摘要）
        if len(self.chunk.content.strip()) <= SUMMARY_LLM_THRESHOLD_CHARS:
            is_leaf = not self.chunk.children
            # 前缀摘要--非叶子节点的原文，叶子节点为空
            self.chunk.prefix_summary = "" if is_leaf else self.chunk.content.strip()
            if is_leaf:
                # 叶子节点摘要-->原文
                self.chunk.summary = self.chunk.content.strip()
            else:
                # 非叶子节点摘要-->原文+子节点的摘要
                parts = [self.chunk.content.strip()] if self.chunk.content.strip() else []
                parts.extend(line for line in lines if line)
                self.chunk.summary = "；".join(parts)
            self._done = True
        else:
            self._thread = threading.Thread(target=self._run_summarize, args=(owner,), daemon=True)
            self._thread.start()

    def _run_summarize(self, owner):
        try:
            actions = owner.execute_prompt('生成节点摘要', self)
            if actions and isinstance(actions[0], dict):
                params = actions[0].get('params', {})
                self.chunk.prefix_summary = params.get('prefix_summary', '')
                self.chunk.summary = params.get('summary', '')
            else:
                GlobalFunction.log('生成节点摘要失败', f'节点 {self.chunk.number} {self.chunk.title} 未获取到有效摘要')
        except Exception as e:
            GlobalFunction.log('节点摘要出错', f'节点 {self.chunk.number} {self.chunk.title}: {e}')
        finally:
            self._done = True

    def Execute(self, owner: StateOwner):
        if self._done:
            owner.remove_state(self)

    def Exit(self, owner):
        pass
            
# 构建pageindex的知识库json知识树
class BuildKnowledgeState(BaseState):
    '''构建json知识树，参数一：源文件路径，参数二，知识库名'''
    stateName = "BuildKnowledgeState"

    def __init__(self, source_path:str , name:str ,**kwargs):
        super().__init__(**kwargs)
        self.path = source_path
        self.name = name
        self._done = False
        self._duration = 1000 * 60 * 30
        self._stage = 'init'
        self._root_summary_done = False
        self._root_node = None

    # 生命周期函数写状态机和owner
    """
    注意java是静态类型语言，使用的时候必须声明才可以编译通过
    python是运行时才检查owner能不能胜任后续的使用，不行-->AttributeError
    鸭子类型-->看起来像鸭子、叫起来像鸭子-->鸭子
    """
    def Enter(self,owner):
        # expanduser针对"~"，翻译为用户目录，resolver绝对值处理，去除.和..这种
        source = Path(self.path).expanduser().resolve()
        if not source.exists():
            GlobalFunction.log("解析错误",f"源文件不存在：{source}")
            self._done = True
            return
        # 建立知识库的目录 parents比如knowledge不存在，设置为true会自动创建
        # exist_ok设置为true，即使目录已经存在也不报错
        output_dir = Path(GlobalFunction.knowledge_path(self.name))
        output_dir.mkdir(parents=True,exist_ok=True)
        # 构建知识库的第一步解析文档
        self._current_sub_state = WordParseState(source_path=self.path,output_root=str(output_dir))
        owner.add_state(self._current_sub_state)
        GlobalFunction.log('解析知识文档','进入知识文档解析的编排线程')
        self._stage = "parse"
        self._output_dir = output_dir #存下来后面要用

    def Execute(self,owner:StateOwner):
        if self._done:
            owner.remove_state(self)
            return
        if self._stage == 'parse':
            if self._current_sub_state not in owner._states:
                self._stage = 'summarize'
                GlobalFunction.log('进入摘要总结阶段','并发处理各层级的摘要中')
                self._chunks = self._current_sub_state._chunks
                self._full_text = self._current_sub_state._full_text
                # 按层级分组，自底向上并发
                self._level_map = {}  # 按照level将节点存起来，dict-->"level":[]
                self._collect_by_level(self._chunks, self._level_map)
                # 先总结子节点
                # sorted函数，第一个参数是可遍历的的结构，reverse true 代表 3,2,1
                self._levels = sorted(self._level_map.keys(), reverse=True)
                self._current_level_idx = 0  # 指针，只想self._levels
                self._pending_summaries = []  # 指针层的状态机实例列表
                self._spawn_current_level(owner)
        elif self._stage == 'summarize':
            # 当前总结摘要的状态机都结束了(pending的都不在states中)
            if not any(s in owner._states for s in self._pending_summaries):
                self._current_level_idx += 1  #拿下一层
                # 不要超过列表边界--跑完每一层
                if self._current_level_idx < len(self._levels):
                    self._spawn_current_level(owner)
                else:
                    self._stage = 'root_summary'
                    GlobalFunction.log('生成根节点摘要','为知识树生成文档级根节点')
                    self._thread = threading.Thread(target=self._run_root_summary, daemon=True, args=(owner,))
                    self._thread.start()

        elif self._stage == 'root_summary':
            if self._root_summary_done:
                self._stage = 'build'
                GlobalFunction.log('建树','建树中（id和落盘）')
                self._thread = threading.Thread(target=self._run_build, daemon=True, args=(owner,))
                self._thread.start()

    def _collect_by_level(self, chunks: list[HeadingChunk], level_map: dict):
        """自上而下，递归收集所有节点，按 level 分组"""
        for chunk in chunks:
            # setdefault 如果key没有对应的value，就返回默认值[]，然后append
            # 如果level已经有对应的value，则直接添加节点
            level_map.setdefault(chunk.level, []).append(chunk)
            if chunk.children:
                self._collect_by_level(chunk.children, level_map)

    def _spawn_current_level(self, owner):
        """为当前层级的所有节点各创建一个 GenSummaryState，实现并发"""
        level = self._levels[self._current_level_idx]  # [3,2,1]第一位是3
        chunks = self._level_map[level]  # 拿到第一层的节点
        self._pending_summaries = []  # 每层节点的状态机列表
        for chunk in chunks:
            gs = GenSummaryState(chunk=chunk)
            self._pending_summaries.append(gs)
            owner.add_state(gs)

    def _chunks_to_json(self,chunks:list[HeadingChunk]):
        """将chunk树-->最终的json树（PageIndex）"""
        jsonChunks = []
        for chunk in chunks:
            node = { #字典字面量
            "node_id" : chunk.node_id,
            "start_index" : chunk.start_index,
            "end_index" : chunk.end_index,
            "summary" : chunk.summary,
            "prefix_summary" : chunk.prefix_summary,
            "use_condition" : chunk.use_condition,
            "level" : chunk.level,
            "number" : chunk.number,
            "title" : chunk.title,
            "children" : self._chunks_to_json(chunk.children),
            "knowledge_name" : self.name
            }
            # 平铺的 但是有索引，谁是谁的子节点。
            jsonChunks.append(node)
        return jsonChunks

    def _run_root_summary(self, owner):
        '为知识树添加一个根节点，方便后续从森林里选出知识树'
        try:
            lines = []
            for chunk in self._chunks:
                lines.append(f"{chunk.number} {chunk.title}")
                if chunk.summary:
                    lines.append(f"  摘要：{chunk.summary}")
            self.root_context = '\n'.join(lines)

            actions = owner.execute_prompt('生成知识树根节点摘要', self)
            params = actions[0].get('params', {}) if actions else {}
            summary = params.get('summary', '')
            use_condition = params.get('use_condition', '')
        except Exception as e:
            GlobalFunction.log('根节点摘要生成失败', f'{e}')
            summary = ''
            use_condition = ''
        finally:
            self._root_node = HeadingChunk(
                level=0,
                number='',
                title=self.name,
                content='',
                children=self._chunks,
                summary=summary,
                use_condition=use_condition,
            )
            GlobalFunction.log('根节点摘要生产完成','根节点生成了全篇摘要')
            self._root_summary_done = True

    def _run_build(self,owner:StateOwner):
        try:
            chunks = [self._root_node] if self._root_node else self._chunks
            if not chunks:
                return
            # DFS分配 node_id，父节点聚合子节点索引
            self._assign_ids_and_indices(chunks=chunks)
            # 最终的知识树
            jsonChunks = self._chunks_to_json(chunks)
            # 落盘json知识树
            GlobalFunction.log('开始落盘知识树',"标志知识树落盘")
            output_path = self._output_dir / '知识树.json'
            output_path.write_text(
                json.dumps(jsonChunks,ensure_ascii=False,indent=2),
                encoding="utf-8"
            )
        except Exception as e:
            GlobalFunction.log('构建失败',f"失败信息：{e}")
        finally:
            # 构建结束
            GlobalFunction.log('知识树落盘成功','知识树落盘成功，注意查收')
            self._done=True
    
    def _assign_ids_and_indices(self, chunks: list[HeadingChunk]):
        """DFS遍历树，分配 node_id，父节点聚合子节点的 start/end_index"""
        counter = [0]

        def walk(nodes: list[HeadingChunk]):
            for chunk in nodes:
                counter[0] += 1
                chunk.node_id = f"{counter[0]:04d}"

                if chunk.children:
                    walk(chunk.children)
                    for child in chunk.children:
                        if child.start_index != -1:
                            if chunk.start_index == -1 or child.start_index < chunk.start_index:
                                chunk.start_index = child.start_index
                        if child.end_index != -1:
                            if chunk.end_index == -1 or child.end_index > chunk.end_index:
                                chunk.end_index = child.end_index

        walk(chunks)
    
    def Exit(self,owner):
        pass

# 根据"方案类型"和"相关的知识树片段"自动生成审查清单
class GenCheckListState(BaseState):
    '''两个阶段：知识树相关性--找出森林里，每棵树上相关的枝条；将知识转换为约束条款--根据每个节点的摘要，总结对应每条的llm审查点'''
    stateName = "GenCheckListState"

    def __init__(self, knowledge_names: list[str], plan_type: str, **kwargs):
        super().__init__(**kwargs)
        self.knowledge_names = knowledge_names
        self.plan_type = plan_type
        self._done = False
        self._duration = 1000*60*30

    def Enter(self, owner: StateOwner):
        output_dir = Path(GlobalFunction.review_path(self.plan_type))
        output_dir.mkdir(parents=True, exist_ok=True)
        self._output_dir = output_dir

        self._thread = threading.Thread(
            target=self._run_generate, args=(owner,), daemon=True
        )
        self._thread.start()

    def _run_generate(self, owner):
        try:
            GlobalFunction.log('开始生成审查清单','生成审查清单线程开启')
            all_trees = []
            for name in self.knowledge_names:
                tree_path = Path(GlobalFunction.knowledge_path(name)) / "知识树.json"
                # extend，是无缝拼接列表，合成新的列表--此处将所有知识树拼在一起，是一个森林
                all_trees.extend(json.loads(tree_path.read_text(encoding="utf-8")))

            selected = self._selected_tree(owner=owner,trees=all_trees)
            self.knowledge_tree = selected
            actions = owner.execute_prompt("生成审查清单", self)
            checklist = actions[0]['params']['checklist']

            output_path = self._output_dir / '审查清单.json'
            output_path.write_text(
                json.dumps(checklist, ensure_ascii=False, indent=2),
                encoding='utf-8'
            )
        except Exception as e:
            GlobalFunction.log('生成审查清单错误', f'错误为{e}')
        finally:
            GlobalFunction.log('成功生成审查清单','注意查收审查清单')
            self._done = True

    def _selected_tree(self,owner,trees:list[dict])->list[dict]:
        """
        此算法为了获得与施工方案相关的各个知识树上分支
        """
        # 逐层导航
        pending = trees
        selected: list[dict] = []

        # 此处是BFS算法，llm自主导航，选出需要的知识树（节点）
        iteration = 0
        while pending and iteration<=15:
            iteration += 1
            lines = []
            for node in pending:
                has = " [含子章节]" if node.get('children') else ""
                uc = node.get('use_condition', '')
                uc_line = f"\n适用条件：{uc}" if uc else ""
                # 转成适合LLM分析的文本 [0001] 1 桩基础 \n 摘要适用条件[含子章节]
                lines.append(f"[{node['node_id']}] {node['number']} {node['title']}\n{node['summary']}  {uc_line}  {has}")
            # llm用的数据-->当前节点的信息
            self.current_nodes = '\n\n'.join(lines) #得到第一层节点信息的字符串，把列表里的每个元素间隔两行

            actions = owner.execute_prompt("导航知识树", self)
            '''
            JSON 对象必须包含：
                - `"action"`: string，固定值 "state:StoreNavigation"
                - `"params"`: object，包含：
                    - `"decisions"`: array，对每个输入节点给出决策，每个元素包含：
                        - `"node_id"`: string，节点唯一标识，与输入一致
                        - `"decision"`: string，"expand" / "select" / "skip"
                        - `"reason"`: string，简述判断依据（20字以内）
            '''
            decisions = actions[0]['params']['decisions']
            # 构造关于节点决策的字面量字典
            decision_map = {d['node_id']: d['decision'] for d in decisions}
            # 下一层(被expand标记)
            next_pending = []
            for node in pending:
                d = decision_map.get(node['node_id'], 'skip')
                # 选中且子节点相关
                if d == 'expand':
                    next_pending.extend(node.get('children', []))
                # 叶子节点被选中
                elif d == 'select':
                    selected.append({
                        "node_id": node['node_id'],
                        "number": node.get('number', ''),
                        "title": node['title'],
                        "summary": node['summary'],
                        "start_index": node.get('start_index', -1),
                        "end_index": node.get('end_index', -1),
                        "knowledge_name": node.get('knowledge_name', ''),
                    })
            # 迭代写法
            pending = next_pending    
        return selected

    def Execute(self,owner:StateOwner):
        if self._done:
            owner.remove_state(self)
            return
    def Exit(self,owner):
        pass
    
# 取"知识库原文"，逐项按照"审查清单"检查"施工方案"合规性
class ReviewPlanState(BaseState):
    stateName = "ReviewPlanState"

    def __init__(self,plan_path:str,plan_type:str,**kwargs):
        super().__init__(**kwargs)
        self.plan_path = plan_path
        self.plan_type = plan_type
        self._stage = 'init'
        self._done = False
        self._duration = 1000*60*15
        
    def Enter(self,owner:StateOwner):
        # 传进来一个施工方案，类似于buildknowledge的思路，先判断在不在
        source = Path(self.plan_path).expanduser().resolve()
        if not source.exists():
            GlobalFunction.log("解析错误",f"源文件不存在：{source}")
            self._done = True
            return
        
        output_dir = Path(GlobalFunction.review_path(self.plan_type))
        output_dir.mkdir(exist_ok=True,parents=True)

        self._current_sub_state = WordParseState(source_path=self.plan_path,output_root=str(output_dir))
        owner.add_state(self._current_sub_state)
        self._stage = "parse"
        self._output_dir = output_dir
        
    def Execute(self,owner:StateOwner):
        if self._done == True:
            owner.remove_state(self)
            return
        if self._stage == 'parse':
            if self._current_sub_state not in owner._states:
                self._stage = 'review'
                self._thread = threading.Thread(
                    # args是按位置一一透传（self是类的实例本身，不是args传的）
                    # 元组和元素的区别(owner就是owner本身，args接受任意可迭代的对象)
                    # 如果owner本身可迭代是不会报错，但是可能数量出现问题；并且也不是我们想要的效果
                    target=self._run_review,args=(owner,),daemon=True  #实例调用，已经绑定了self。类名.方法，才需要self参数也写
                )
                self._thread.start()
            
    def Exit(self,owner):
        pass
    
    '''
    daemon线程的异常不会外抛，直接死掉。使用try/except
    '''
    def _run_review(self, owner:StateOwner):
        '''
        拿到审查清单+对应原文-->llm开始分析-->输出结果
        '''
        try:
            GlobalFunction.log('开始依据审查清单逐项审查','正在逐条审查方案合规性')
            # wordparsestate挂载的--方案原文
            self.plan_text = self._current_sub_state._full_text
            # 拿到审查清单
            checklist_path = Path(GlobalFunction.review_path(self.plan_type)) / "审查清单.json"
            checklist:list[dict] = json.loads(checklist_path.read_text(encoding='utf-8'))
            
            # 条规原文
            full_texts= {} 
            result:list[dict] = []  # 最后的输出结果
            
            for check in checklist:
                self.check_item = check['check_item']
                self.check_method = check['check_method']
                self.level = check['level']
                
                # 取条款原文
                name = check.get('knowledge_name', '')
                start = check.get('start_index', -1)
                end = check.get('end_index', -1)
                
                if name and start != -1 and end != -1:
                    if name not in full_texts: # 懒加载，拿到对应知识树的全文
                        dir_path = Path(GlobalFunction.knowledge_path(name))
                        files = list(dir_path.glob("**/*_全文.txt"))
                        full_texts[name] = files[0].read_text(encoding='utf-8') if files else ''
                    self.clause_text = full_texts[name][start:end]
                else:
                    self.clause_text = check['check_item']  # 没有拿到文本--鲁棒性采用审查点替代
                
                actions = owner.execute_prompt('逐条合规检查', self)
                params: dict= actions[0]['params']
                '''
                {
                    "action": "state:StoreReviewResult",
                    "params": {
                        "clause_number": "3.2.1",
                        "result": "合规",
                        "reason": "方案明确规定了试桩数量、目的和清孔工艺要求",
                        "evidence": "正式施工前选取3根工程桩作为试桩，通过试桩确定钻机转速、泥浆比重等工艺参数。成孔后采用反循环清孔工艺，清孔后沉渣厚度不大于50mm。"
                    }
                }
                '''
                record = {**check, **params}
                result.append(record)
            
            result_path = self._output_dir / '审查结果.json'
            result_path.write_text(json.dumps(result, indent=2, ensure_ascii=False), encoding='utf-8')
        except Exception as e:
            GlobalFunction.log('审查线程错误', f'错误为{e}')
        finally:
            GlobalFunction.log('合规性检查完成','可以查看审查结果')
            self._done = True

# 人工复查环节（涉及open webui）
class HumanReviewState(BaseState):
    def __init__(self,**kwargs):
        super().__init__(**kwargs)
    def Enter(self, owner):
        pass
    def Execute(self, owner):
        pass
    def Exit(self, entity):
        pass
    
# 生成报告（提示词--审查报告）
class GenReportState(BaseState):
    def __init__(self, prompt:str, data:list, plan_type:str, **kwargs):
        super().__init__(**kwargs)
        self.prompt = prompt
        self.data = data
        self.plan_type = plan_type
        self._duration = 1000*60*10
        self._done = False
        self._thread = None

    def Enter(self, owner):
        self.data = json.dumps(self.data, ensure_ascii=False, indent=2)
        output_path = Path(GlobalFunction.review_path(self.plan_type))
        output_path.mkdir(parents=True, exist_ok=True)
        self._output_dir = output_path
        self._thread = threading.Thread(target=self._gen_report, args=(owner,), daemon=True)
        self._thread.start()

    def Execute(self, owner:StateOwner):
        if self._done:
            owner.remove_state(self)

    def Exit(self, entity):
        pass

    def _gen_report(self, owner:StateOwner):
        try:
            GlobalFunction.log('开始生成报告','进入报告生成线程')
            actions = owner.execute_prompt(self.prompt, self)
            report_md = actions[0]['params']['report']
            if not report_md:
                raise ValueError('生成报告的内容为空')

            md_path = self._output_dir / '审查报告.md'
            md_path.write_text(report_md, encoding='utf-8')
            GlobalFunction.log('报告md已保存', str(md_path))

            docx_path = self._output_dir / '审查报告.docx'
            self._md_to_docx(report_md, str(docx_path))
            GlobalFunction.log('报告docx已生成', str(docx_path))
        except Exception as e:
            GlobalFunction.log('生成报告线程出错', f'错误为:{e}')
        finally:
            GlobalFunction.log('成功生成报告', '注意查收审查报告')
            self._done = True

    def _md_to_docx(self, md_text: str, output_path: str):
        """将 markdown 文本转为 docx，处理标题、表格、加粗、列表"""
        lines = md_text.strip().split('\n')
        doc = Document()
        doc.styles['Normal'].font.size = Pt(12)
        doc.styles['Normal'].font.name = '宋体'

        i = 0
        while i < len(lines):
            line = lines[i]

            if line.strip().startswith('|') and line.strip().endswith('|'):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                    table_lines.append(lines[i])
                    i += 1
                rows = [[c.strip() for c in ln.strip().split('|')[1:-1]] for ln in table_lines]
                if len(rows) >= 2:
                    header = rows[0]
                    sep_re = re.compile(r'^[-:]{3,}$')
                    data_start = 2 if (len(rows) > 1 and all(sep_re.match(c) for c in rows[1])) else 1
                    table = doc.add_table(rows=1 + len(rows) - data_start, cols=len(header))
                    table.style = 'Light Grid Accent 1'
                    for j, cell_text in enumerate(header):
                        table.rows[0].cells[j].text = ''
                        self._write_cell(table.rows[0].cells[j].paragraphs[0], cell_text)
                    for row_idx in range(data_start, len(rows)):
                        for j, cell_text in enumerate(rows[row_idx]):
                            if j < len(header):
                                table.rows[1 + row_idx - data_start].cells[j].text = ''
                                self._write_cell(table.rows[1 + row_idx - data_start].cells[j].paragraphs[0], cell_text)
                continue

            heading_match = re.match(r'^(#{1,3})\s+(.+)', line)
            if heading_match:
                level = len(heading_match.group(1))
                text = re.sub(r'\*\*([^*]+)\*\*', r'\1', heading_match.group(2))
                doc.add_heading(text, level=level)
                i += 1
                continue

            list_match = re.match(r'^(\s*)-\s+(.+)', line)
            if list_match:
                p = doc.add_paragraph(style='List Bullet')
                self._write_cell(p, list_match.group(2))
                i += 1
                continue

            ordered_match = re.match(r'^(\s*)\d+\.\s+(.+)', line)
            if ordered_match:
                p = doc.add_paragraph(style='List Number')
                self._write_cell(p, ordered_match.group(2))
                i += 1
                continue

            if not line.strip() or re.match(r'^[-*_]{3,}\s*$', line):
                i += 1
                continue

            p = doc.add_paragraph()
            self._write_cell(p, line)
            i += 1

        doc.save(output_path)

    @staticmethod
    def _write_cell(paragraph, text: str):
        """写入段落，解析 **bold** 标记"""
        for part in re.split(r'(\*\*[^*]+\*\*)', text):
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)
